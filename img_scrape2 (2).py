from tkinter import *
import datetime
import urllib.request
import csv
from docx import Document
from docx.shared import Inches
from docx.shared import Pt



root=Tk()
root.geometry("1200x1200")
root.title("Cake-Shopping Site")

TopFrame=Frame(root,width=1350,height=10,bd=14,relief='raise')
TopFrame.pack(side=TOP)

BottomFrame=Frame(root,width=350,height=50,bd=14,relief='raise')
BottomFrame.pack(side=BOTTOM)
LeftMidFrame=Frame(root,width=1050,height=1150,bd=14,relief="raise")
LeftMidFrame.pack(side=LEFT)

lbtTitle=Label(TopFrame,font=('arial',24,'bold'),text="Delicious Bakeries",bd=10,width=45,justify='center')
lbtTitle.grid(row=1,column=1)


global var8,var7,var4,var5,var6,var12;
var2=IntVar()
var3=IntVar()
var4=StringVar()#link
var5=StringVar()#title
var6=StringVar()
var12=StringVar()
var7=StringVar()
#var8=lbladd_input.get()
var9=StringVar()
var8=StringVar()
var10=IntVar()
var11=StringVar()

#var1.set("0")
#var2.set("0")
#var3.set("0")




lbtProductID=Label(LeftMidFrame,font=('arial',18,'bold'),text="Product ID",bd=1,width=12,justify='left')
lbtProductID.grid(row=0,column=0)
lbtProductID_input=Entry(LeftMidFrame,font=('arial',18,'bold'),textvariable=var2,bd=1,width=12,justify='left')
lbtProductID_input.grid(row=0,column=1)


lblPrice=Label(LeftMidFrame,font=('arial',18,'bold'),text="Price",bd=10,width=30)#var4 cake price
lblPrice.grid(row=1,column=0)
lblPrice_input=Label(LeftMidFrame,font=('arial',14,'bold'),textvariable=var4,bd=10,width=55,relief="sunken")#filled on pressing button SET
lblPrice_input.grid(row=1,column=1)

lbltype=Label(LeftMidFrame,font=('arial',18,'bold'),text="Cake Type",bd=10,width=30)#var6 cake type
lbltype.grid(row=2,column=0)
lbltype_input=Label(LeftMidFrame,font=('arial',14,'bold'),textvariable=var6,bd=10,width=55,relief="sunken")#filled on pressing button SET
lbltype_input.grid(row=2,column=1)

lblPrTitle=Label(LeftMidFrame,font=('arial',18,'bold'),text="Product Name",bd=10,width=20)#var5 cake title
lblPrTitle.grid(row=3,column=0)
lblPrTitle_input=Label(LeftMidFrame,font=('arial',14,'bold'),textvariable=var5,bd=10,width=30,relief="sunken")#filled on pressing button SET
lblPrTitle_input.grid(row=3,column=1)

lblName=Label(LeftMidFrame,font=('arial',18,'bold'),text="Customer Name",bd=10,width=20)
lblName.grid(row=4,column=0)
lblName_input=Entry(LeftMidFrame,font=('arial',14,'bold'),textvariable=var7,bd=10,width=30,relief="sunken")#filled on pressing button SET
lblName_input.grid(row=4,column=1)

lbladd=Label(LeftMidFrame,font=('arial',18,'bold'),text="Customer Address",bd=10,width=20)
lbladd.grid(row=5,column=0)
lbladd_input=Entry(LeftMidFrame,font=('arial',14,'bold'),textvariable=var8,bd=10,width=30,relief="sunken")#filled on pressing button SET
lbladd_input.grid(row=5,column=1)

lbltime=Label(LeftMidFrame,font=('arial',18,'bold'),text="Current date & time",bd=10,width=20)
lbltime.grid(row=6,column=0)
lbltime_input=Label(LeftMidFrame,font=('arial',14,'bold'),textvariable=var9,bd=10,width=30,relief="sunken")#filled on pressing button SET
lbltime_input.grid(row=6,column=1)

lbldis=Label(LeftMidFrame,font=('arial',18,'bold'),text="Approximate distance from us (in km)",bd=10,width=40)
lbldis.grid(row=7,column=0)
lbldist_input=Entry(LeftMidFrame,font=('arial',14,'bold'),textvariable=var10,bd=10,width=30,relief="sunken")#filled on pressing button SET
lbldist_input.grid(row=7,column=1)


lbldel=Label(LeftMidFrame,font=('arial',18,'bold'),text="Delivery time :",bd=10,width=30)
lbldel.grid(row=8,column=0)
lbldel_input=Label(LeftMidFrame,font=('arial',14,'bold'),textvariable=var11,bd=10,width=30,relief="sunken")#filled on pressing button SET
lbldel_input.grid(row=8,column=1)

lblmail=Label(LeftMidFrame,font=('arial',18,'bold'),text="Customer E-Mail ID",bd=10,width=20)
lblmail.grid(row=9,column=0)
lblmail_input=Entry(LeftMidFrame,font=('arial',14,'bold'),textvariable=var12,bd=10,width=30,relief="sunken")#filled on pressing button SET
lblmail_input.grid(row=9,column=1)

'''cmbProductID=ttk.Combobox(BottomFrame,textvariable=var10,state='readonly',font={'arial',28,'bold'},width=13)
cmbProductID['value']=("0-25 km","25-50 km","50-75 km","75-100 km")
#cmbProductID['value']=(2,3,4,5)
cmbProductID.current("0 km")
cmbProductID.grid(row=7,column=1)

'''







i=0
import urllib.request

def set_val():
    with open('scraped.csv')as f:
        save=csv.reader(f)
        for row in save:
            print(row[i])
            print(var2.get())
            if str(var2.get())==str(row[i]):
                var4.set(row[i+4])
                var5.set(row[i+2])
                var6.set(row[i+3])
                var9.set(datetime.datetime.now())
                ''' urllib.request.urlretrieve(row[i+1],"/home/ankita/Desktop/Python files/img")
                from PIL import Image
                img = Image.open("/home/ankita/Desktop/Python files/img")
                img.save('img','gif')
                photo=PhotoImage(file=("/home/ankita/Desktop/Python files/img"))


                cv=Canvas()
                cv.pack(side="left",fill="both",expand="yes")
                cv.create_image(100,200,image=photo,anchor='nw')'''
                break
            else:
                print("\nMoving onto next")
                print(lblmail_input.get())

            if var10.get()<25:
                var11.set("1 hr")
            elif var10.get()<50 and var10.get()>25:
                var11.set("2 hrs.")
            elif var10.get()<75 and var10.get()>50:
                var11.set("3 hrs")
            elif var10.get()<100 and var10.get()>75:
                var11.set("4 hrs.")
            else:
                var11.set("Sorry This order cannot be placed.It's out of reach!!")




def generate_doc():
    #v1 = professor_name_entry.get()
    #v2 = professor_desg_entry.get()
    #v3 = professor_colg_entry.get()
    var8=lbladd_input.get()
    
    var7=lblName_input.get()
   # var5=lblPrTitle_input.get()
    #var4=lblPrice_input.get()
   # var6=lbltype_input.get()
    document = Document("noc.docx")

    p = document.add_paragraph()
    p.style = document.styles['Normal']
    font = p.style.font
    
    
    font.name = 'Symbola'
    font.size = Pt(16)
    #p.add_run('\nCake Bill, \n').bold = True
    #p.add_run(var5).bold=True
    p.add_run("\n\n                              Customer Bill\n")
    p.add_run(" *******************************************************")
    p.add_run('\n\n')
    #p.add_run(var7).bold=True
    p.add_run("         ")
    p.add_run('Customer Name:').bold = True
    p.add_run(var7).bold=False
    p.add_run('\n\n')
    p.add_run("         ")
    p.add_run('Customer Address:').bold=True
    p.add_run(var8).bold=False
    p.add_run('\n\n')
    p.add_run("         ")
    p.add_run('Ordered Cake:').bold=True
    p.add_run(var5.get()).bold=False
    p.add_run('\n\n')
    p.add_run("         ")
    p.add_run('Type:').bold=True
    p.add_run(var6.get()).bold=False
    p.add_run('\n\n')
    p.add_run("         ")
    p.add_run('Cake Price:').bold=True
    p.add_run(var4.get()).bold=False
    p.add_run('\n\n\n')
    #r=p.add_run()
    #r.add_picture(img,width=1200000,height=1200000)
    p.add_run("\n\n\n\n\n\n")
    p.add_run("********************************************************")
    #font.size=Pt(10)
    p.add_run("\nContact us:")
    p.add_run("\nMail: bakingocakes@gmail.com                                   Phone no:8765432109")
    p.add_run("\n                          Thank You for visiting.")
    document.save('sample2.docx')
    print('notreach')
    #o=str(var12.get())
    o=lblmail_input.get()
    print('reach')
    print(o)
    import smtplib
    smtpObj=smtplib.SMTP_SSL('smtp.gmail.com',465)
    smtpObj.ehlo()
    smtpObj.login('ankidas44@gmail.com','DinuChacha5*')
    msg='Subject:Your order has been placed!!'
    smtpObj.sendmail('ankidas44@gmail.com',o,msg)
    smtpObj.quit()
btn2 = Button(BottomFrame, text='Generate bill', command=generate_doc)
btn2.grid(row=10,column=1)    

btnTotal=Button(BottomFrame,font=('arial',18,'bold'),text='Total',bd=2,command=set_val)
btnTotal.grid(row=10,column=0)




'''img.save('img','gif')
photo=PhotoImage(file=("/home/ankita/Desktop/Python files/img"))

cv=Canvas()
cv.pack(side="left",fill="both",expand="yes")
cv.create_image(100,200,image=photo,anchor='nw')

'''
'''Email Sending!'''

#var12=lblmail_input.get()


